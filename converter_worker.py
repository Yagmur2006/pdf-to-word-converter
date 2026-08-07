#!/usr/bin/env python3
"""
pdf2word - CPU-bound worker module.

Everything in this module is designed to run inside a *separate process*
(``ProcessPoolExecutor``) so that the Flask event loop / request threads are
never blocked by PDF parsing.

Parallelism strategy
--------------------
``pdf2docx`` ships a ``multi_processing=True`` flag, but its implementation
serialises intermediate results to ``pages-{i}.json`` **in the current working
directory**, which race-conditions across concurrent jobs on a web server.

Instead we drive the same four-step pipeline manually:

    1. ``load_pages()``      -> open the PDF                (per worker)
    2. ``parse_document()``  -> section / margin analysis   (per worker)
    3. ``parse_pages()``     -> layout parsing              (PARALLEL, chunked)
    4. ``make_docx()``       -> docx assembly               (single worker)

Step 3 is the expensive one (image clipping, shape/table detection), so it is
split into page chunks that are parsed concurrently across CPU cores. Each
chunk is serialised to its own JSON file inside a job-private temp directory
(no shared filenames -> no races), then a single "build" task deserialises the
chunks in order and emits the final ``.docx``.
"""

from __future__ import annotations

import gc
import json
import logging
import math
import os
from typing import Any

# ---------------------------------------------------------------------------
# Quality profiles - real pdf2docx tuning knobs
# ---------------------------------------------------------------------------
# `clip_image_res_ratio` is by far the heaviest setting: it is the resolution
# multiplier (relative to 72dpi) used when rasterising vector graphics / clipped
# page regions. Lowering it from 4.0 -> 1.5 cuts both RAM and time dramatically
# on image-heavy documents (pixel count scales with the square of the ratio).
BASE_SETTINGS: dict[str, Any] = {
    "debug": False,
    "ocr": 0,
    "ignore_page_error": True,   # a single bad page must not kill the job
    "raw_exceptions": False,
    "multi_processing": False,   # we do our own, race-free parallelism
    "cpu_count": 0,
    "list_not_table": True,
}

QUALITY_PROFILES: dict[str, dict[str, Any]] = {
    # Default and only profile exposed to the UI. Text extraction is never
    # touched by any of these knobs - only vector-shape / table / image
    # analysis is toned down. Measured on dense math-formula / scanned-page
    # PDFs (thousands of tiny vector paths per page, the worst case for
    # pdf2docx): `min_svg_w`/`min_svg_h` alone is responsible for a 2.5-3.6x
    # speedup, because it lets pdf2docx discard microscopic decorative shapes
    # *before* the O(n^2)-ish adjacent-shape merge pass instead of after.
    "fast": {
        "clip_image_res_ratio": 1.5,
        "extract_stream_table": False,
        "parse_stream_table": False,
        "parse_lattice_table": False,
        "min_svg_w": 12.0,
        "min_svg_h": 12.0,
        "min_svg_gap_dx": 20.0,
        "min_svg_gap_dy": 4.0,
        "shape_min_dimension": 5.0,
    },
    # Kept for API completeness (e.g. ?quality=balanced) - not reachable from
    # the UI. Closer to stock pdf2docx behaviour for callers who need it.
    "balanced": {
        "clip_image_res_ratio": 2.0,
        "extract_stream_table": False,
        "parse_stream_table": True,
        "parse_lattice_table": True,
        "min_svg_w": 3.0,
        "min_svg_h": 3.0,
        "shape_min_dimension": 2.0,
    },
    # Stock pdf2docx behaviour, maximum fidelity.
    "accurate": {
        "clip_image_res_ratio": 4.0,
        "extract_stream_table": True,
        "parse_stream_table": True,
        "parse_lattice_table": True,
    },
}

DEFAULT_QUALITY = "fast"


def build_settings(quality: str = DEFAULT_QUALITY) -> dict[str, Any]:
    """Return the *partial* override dict for a quality profile."""
    overrides = dict(BASE_SETTINGS)
    overrides.update(QUALITY_PROFILES.get(quality, QUALITY_PROFILES[DEFAULT_QUALITY]))
    return overrides


# ---------------------------------------------------------------------------
# Process pool initializer
# ---------------------------------------------------------------------------
def _die_with_parent() -> None:
    """Ask the kernel to SIGTERM this worker if the web process disappears.

    Without this, a hard-killed Flask parent leaves orphaned conversion
    workers holding hundreds of MB of RAM.
    """
    try:
        import ctypes
        import signal

        libc = ctypes.CDLL("libc.so.6", use_errno=True)
        PR_SET_PDEATHSIG = 1
        libc.prctl(PR_SET_PDEATHSIG, signal.SIGTERM)

        # Guard the race where the parent died before prctl was installed.
        if os.getppid() == 1:
            os._exit(0)
    except Exception:  # non-Linux / restricted environments
        pass


def init_worker() -> None:
    """Runs once per pool worker process.

    Pins numeric libraries to a single thread. OpenCV / NumPy (pulled in by
    pdf2docx) otherwise spin up one thread *per core inside every worker*,
    which oversubscribes the CPU and slows the whole pool down.
    """
    _die_with_parent()

    for var in (
        "OMP_NUM_THREADS",
        "OPENBLAS_NUM_THREADS",
        "MKL_NUM_THREADS",
        "NUMEXPR_NUM_THREADS",
        "VECLIB_MAXIMUM_THREADS",
    ):
        os.environ.setdefault(var, "1")

    logging.getLogger().setLevel(logging.ERROR)
    for name in ("pdf2docx", "fitz", "PIL"):
        logging.getLogger(name).setLevel(logging.ERROR)

    try:  # keep OpenCV from spawning its own thread pool too
        import cv2  # noqa: WPS433 (import inside function is intentional)

        cv2.setNumThreads(1)
    except Exception:  # pragma: no cover - cv2 always present via pdf2docx
        pass


def preload() -> None:
    """Import the heavy conversion stack in the *parent* process.

    Pool workers are forked, so anything imported here is inherited by every
    worker for free - no per-worker import cost, and no multi-second stall on
    the first user request.
    """
    import fitz  # noqa: F401
    from pdf2docx.converter import Converter  # noqa: F401


def warmup() -> str:
    """Trivial task that forces an idle pool worker to actually be created."""
    import pdf2docx  # noqa: F401

    return "ready"


# ---------------------------------------------------------------------------
# PDF probing (cheap, safe to call from the web process)
# ---------------------------------------------------------------------------
def probe_pdf(pdf_path: str) -> dict[str, Any]:
    """Open a PDF and return basic metadata without parsing any layout.

    Raises:
        ValueError: if the file is not a readable / decryptable PDF.
    """
    import fitz  # PyMuPDF, already a pdf2docx dependency

    doc = None
    try:
        doc = fitz.Document(pdf_path)
        if doc.needs_pass:
            raise ValueError("Password protected PDFs are not supported.")
        pages = doc.page_count
        if pages <= 0:
            raise ValueError("The PDF contains no pages.")
        return {"pages": pages}
    except ValueError:
        raise
    except Exception as exc:
        # Never surface the internal temp path from the underlying error.
        raise ValueError(
            "The file is not a readable PDF (it may be corrupted or renamed "
            "from another format)."
        ) from exc
    finally:
        if doc is not None:
            try:
                doc.close()
            except Exception:
                pass


def plan_chunks(
    page_count: int,
    pages_per_chunk: int,
    max_chunks: int,
    min_chunks: int = 1,
) -> list[list[int]]:
    """Split ``page_count`` pages into balanced, contiguous index chunks.

    ``min_chunks`` lets the caller guarantee that a single small-but-expensive
    document still saturates every CPU core, while a large batch (where
    cross-file parallelism already does that) can fall back to one chunk per
    file to avoid redundant per-chunk document analysis.
    """
    if page_count <= 0:
        return []

    n_chunks = max(1, math.ceil(page_count / max(1, pages_per_chunk)))
    n_chunks = max(n_chunks, min_chunks)
    n_chunks = min(n_chunks, max_chunks, page_count)
    base, remainder = divmod(page_count, n_chunks)

    chunks: list[list[int]] = []
    cursor = 0
    for i in range(n_chunks):
        size = base + (1 if i < remainder else 0)
        if size == 0:
            continue
        chunks.append(list(range(cursor, cursor + size)))
        cursor += size
    return chunks


# ---------------------------------------------------------------------------
# Stage 1 (parallel): parse a chunk of pages -> JSON
# ---------------------------------------------------------------------------
def parse_chunk(
    pdf_path: str,
    page_indexes: list[int],
    out_json: str,
    overrides: dict[str, Any],
) -> str:
    """Parse ``page_indexes`` of ``pdf_path`` and serialise them to ``out_json``.

    Executed in a pool worker. Returns the JSON path so the parent can restore
    chunks in the right order.
    """
    from pdf2docx.converter import Converter

    cv = None
    try:
        cv = Converter(pdf_path)
        settings = cv.default_settings
        settings.update(overrides)

        cv.load_pages()

        # Only parse the pages belonging to this chunk. Everything else is
        # marked `skip_parsing` so no layout work (and no image clipping)
        # happens for pages owned by sibling workers.
        wanted = set(page_indexes)
        for page in cv.pages:
            page.skip_parsing = page.id not in wanted

        cv.parse_document(**settings)
        cv.parse_pages(**settings)
        cv.serialize(out_json)
        return out_json
    except MemoryError as exc:
        raise RuntimeError(
            "Ran out of memory while parsing this PDF. Try the 'Fast' mode "
            "or split the document into smaller files."
        ) from exc
    finally:
        if cv is not None:
            try:
                cv.close()
            except Exception:
                pass
        gc.collect()


# ---------------------------------------------------------------------------
# Stage 2 (parallel, one call per build group): chunks -> partial .docx
# ---------------------------------------------------------------------------
# `make_docx()` is inherently serial *within one Converter/Document instance*,
# which is why a single huge file used to spend its entire assembly phase on
# one CPU core. To parallelise it, the caller splits a file's parsed chunks
# into up to `PROCESS_WORKERS` groups and calls this function once per group,
# in parallel pool workers - each restoring only ITS OWN pages and writing a
# small partial ``.docx``. `merge_docx_parts()` then stitches the partials
# back together in order. For a single-group file (small documents), the
# caller passes all chunks at once and this *is* already the final document.
def build_docx(
    pdf_path: str,
    json_files: list[str],
    docx_path: str,
    overrides: dict[str, Any],
) -> str:
    """Restore the given parsed page chunks and write them to ``docx_path``.

    When ``json_files`` covers every chunk of a document, this produces the
    final ``.docx`` directly. When it covers only a subset (a "build group"),
    the result is a partial document meant to be combined by
    ``merge_docx_parts()``.
    """
    from pdf2docx.converter import Converter

    cv = None
    try:
        cv = Converter(pdf_path)
        settings = cv.default_settings
        settings.update(overrides)

        # `restore()` seeds the page container from `page_cnt` on first call,
        # then slots every page back by its own id -> original order is kept
        # regardless of which worker produced which chunk.
        restored = False
        for path in json_files:
            if not os.path.exists(path):
                continue
            with open(path, "r", encoding="utf-8") as handle:
                cv.restore(json.load(handle))
            restored = True

        if not restored:
            raise RuntimeError("No page data was produced for this document.")

        cv.make_docx(docx_path, **settings)
        return docx_path
    except MemoryError as exc:
        raise RuntimeError(
            "Ran out of memory while assembling the Word document. "
            "Try the 'Fast' mode or a smaller file."
        ) from exc
    finally:
        if cv is not None:
            try:
                cv.close()
            except Exception:
                pass
        # Intermediate JSON can be hundreds of MB for big documents - drop it
        # as soon as the docx exists instead of waiting for the job sweeper.
        for path in json_files:
            try:
                os.remove(path)
            except OSError:
                pass
        gc.collect()


# ---------------------------------------------------------------------------
# Stage 3 (single worker, fast): merge partial .docx files -> final .docx
# ---------------------------------------------------------------------------
def merge_docx_parts(part_paths: list[str], out_path: str) -> str:
    """Concatenate build-group partials into one ``.docx``, in order.

    Uses ``docxcompose`` to copy each partial's paragraphs, tables, styles and
    *image relationships* into a single master document - this is pure XML /
    zip-part manipulation, not page re-rendering, so it is fast even for
    documents with many pages or images.

    IMPORTANT: ``docxcompose.Composer.append()`` does **not** insert a page
    break at the seam between two documents (verified empirically - the last
    paragraph of part N and the first paragraph of part N+1 land in the same
    section with no break between them). Each partial's *internal* page
    breaks (one per original PDF page, added by pdf2docx itself) are
    preserved correctly; only the *boundary* between partials needs an
    explicit, manually inserted page break, added here before every merge.
    """
    from docx import Document
    from docxcompose.composer import Composer

    if not part_paths:
        raise RuntimeError("No document parts to merge.")

    if len(part_paths) == 1:
        if os.path.abspath(part_paths[0]) != os.path.abspath(out_path):
            shutil.copyfile(part_paths[0], out_path)
        return out_path

    master = Document(part_paths[0])
    composer = Composer(master)
    try:
        for part_path in part_paths[1:]:
            # Force a real page break at the seam (see note above) - a plain
            # run-level break, not a section break, so it cannot disturb the
            # page size/margins the next part's own pages establish.
            master.add_page_break()
            composer.append(Document(part_path))
        composer.save(out_path)
    finally:
        gc.collect()

    return out_path
